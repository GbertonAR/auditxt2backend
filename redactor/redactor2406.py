from fastapi import APIRouter, HTTPException, Query
from pydantic import BaseModel
# CAMBIO CLAVE: Importar AsyncAzureOpenAI
from openai import AsyncAzureOpenAI
from Backend_app.config import settings
import logging
import os
from docx import Document

# Configurar logger
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

router = APIRouter()

# Configuración de Azure OpenAI
AZURE_OPENAI_KEY = settings.azure_openai_key
AZURE_OPENAI_ENDPOINT = settings.azure_openai_endpoint
AZURE_DEPLOYMENT_NAME = settings.azure_openai_deployment

logger.info(f"🔑 Cargando configuración Azure OpenAI")
logger.info(f"Endpoint: {AZURE_OPENAI_ENDPOINT}")
logger.info(f"Deployment: {AZURE_DEPLOYMENT_NAME}")

# CAMBIO CLAVE: Instanciar AsyncAzureOpenAI en lugar de AzureOpenAI
client = AsyncAzureOpenAI(
    api_key=AZURE_OPENAI_KEY,
    api_version="2024-02-15-preview", # Asegúrate de que esta API version sea compatible con tu despliegue de Azure
    azure_endpoint=AZURE_OPENAI_ENDPOINT
)

# Modelo de entrada
class RedaccionRequest(BaseModel):
    titulo: str
    contenido: str
    tono: str = "institucional"
    audiencia: str = "público general"
    
#Modelo de salida
class Articulo(BaseModel):
    titulo: str
    contenido: str
    autor: str
    
# Función para generar contenido con Azure OpenAI
async def generar_contenido_ia(prompt: str, tono: str, audiencia: str) -> str:
    system_prompt = (
        f"Eres un redactor oficial del departamento de prensa de un organismo estatal. "
        f"Redacta el contenido solicitado con tono {tono}, dirigido al {audiencia}. "
        f"Asegúrate de que el mensaje sea claro, institucional, empático y socialmente responsable."
    )

    try:
        logger.info("📨 Enviando solicitud a Azure OpenAI...")
        # CAMBIO CLAVE: El método create del cliente asíncrono ya es awaitable.
        # No necesitas pasar deployment_id como un argumento separado aquí si ya está en 'model'.
        response = await client.chat.completions.create(
            model=settings.azure_openai_deployment, # Utiliza el nombre del deployment aquí
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=800
        )
        resultado = response.choices[0].message.content
        logger.info("✅ Contenido generado correctamente")
        return resultado
    except Exception as e:
        logger.error(f"❌ Error al generar contenido con OpenAI: {str(e)}")
        raise # Re-lanza la excepción para que FastAPI la capture y devuelva un 500

# Ruta de la API
@router.post("/generar")
async def generar_contenido(data: RedaccionRequest):
    print(f"📥 Solicitud recibida: {data}")
    try:
        logger.info(f"📥 Solicitud recibida: {data}")
        resultado = await generar_contenido_ia(data.contenido, data.tono, data.audiencia)
        return {"titulo": data.titulo, "resultado": resultado}
    except Exception as e:
        logger.error(f"❌ Error en endpoint /generar: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error al generar contenido: {str(e)}")
    
@router.post("/guardar-articulo")
async def descargar_articulo(data: Articulo, formato: str = Query("txt", enum=["txt", "docx"])):
    try:
        os.makedirs("articulos", exist_ok=True)

        filename = f"{data.titulo}.{formato}"
        filepath = os.path.join("articulos", filename)

        if formato == "txt":
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(f"Título: {data.titulo}\n")
                f.write(f"Autor: {data.autor}\n\n")
                f.write(data.contenido)
        else:
            doc = Document()
            doc.add_heading(data.titulo, 0)
            doc.add_paragraph(f"Autor: {data.autor}")
            doc.add_paragraph("")
            doc.add_paragraph(data.contenido)
            doc.save(filepath)

        return FileResponse(filepath, media_type="application/octet-stream", filename=filename)

    except Exception as e:
        return {"error": str(e)}   