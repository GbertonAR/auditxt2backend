from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import FileResponse
from pydantic import BaseModel
from openai import AsyncAzureOpenAI
from Backend_app.config import settings
import logging
import os
import asyncio
from docx import Document

# ───────────────────────────── LOGGING ─────────────────────────────
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# ───────────────────────────── ROUTER ─────────────────────────────
router = APIRouter(tags=["Redactor"])

# ────────────────────── CONFIGURACIÓN OPENAI ──────────────────────
AZURE_OPENAI_KEY = settings.azure_openai_key
AZURE_OPENAI_ENDPOINT = settings.azure_openai_endpoint
AZURE_DEPLOYMENT_NAME = settings.azure_openai_deployment

logger.info("🔐 Cargando configuración Azure OpenAI")
logger.info(f"📍 Endpoint: {AZURE_OPENAI_ENDPOINT}")
logger.info(f"🚀 Deployment: {AZURE_DEPLOYMENT_NAME}")

client = AsyncAzureOpenAI(
    api_key=AZURE_OPENAI_KEY,
    api_version="2024-02-15-preview",
    azure_endpoint=AZURE_OPENAI_ENDPOINT
)

# ────────────────────── MODELOS DE DATOS ──────────────────────
class RedaccionRequest(BaseModel):
    titulo: str
    contenido: str
    tono: str = "institucional"
    audiencia: str = "público general"

class Articulo(BaseModel):
    titulo: str
    contenido: str
    autor: str

# ─────────────── GENERACIÓN DE CONTENIDO IA ───────────────
async def generar_contenido_ia(prompt: str, tono: str, audiencia: str) -> str:
    system_prompt = (
        f"Eres un redactor oficial del departamento de prensa de un organismo estatal. "
        f"Redacta el contenido solicitado con tono {tono}, dirigido al {audiencia}. "
        f"Asegúrate de que el mensaje sea claro, institucional, empático y socialmente responsable."
    )

    for intento in range(3):
        try:
            logger.info("📨 Enviando solicitud a Azure OpenAI...")
            response = await client.chat.completions.create(
                model=AZURE_DEPLOYMENT_NAME,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=800
            )
            resultado = response.choices[0].message.content
            logger.info("✅ Contenido generado correctamente")
            return resultado

        except Exception as e:
            logger.warning(f"⚠️ Error al intentar generar contenido (intento {intento+1}): {e}")
            if intento < 2:
                await asyncio.sleep(2 ** intento)  # backoff exponencial
            else:
                logger.error("❌ Fallo persistente al generar contenido")
                raise

# ──────────────── ENDPOINT: Generar contenido ────────────────
@router.post("/generar", summary="Generar contenido redactado con IA")
async def generar_contenido(data: RedaccionRequest):
    logger.info(f"📥 Solicitud recibida: {data}")
    try:
        resultado = await generar_contenido_ia(data.contenido, data.tono, data.audiencia)
        return {"titulo": data.titulo, "resultado": resultado}
    except Exception as e:
        logger.error(f"❌ Error en /generar: {e}")
        raise HTTPException(status_code=500, detail=f"Error al generar contenido: {str(e)}")

# ──────────────── ENDPOINT: Descargar artículo ────────────────
@router.post("/guardar-articulo", summary="Guardar artículo generado en .txt o .docx")
async def descargar_articulo(data: Articulo, formato: str = Query("txt", enum=["txt", "docx"])):
    try:
        output_dir = "/tmp"  # ubicación segura en Azure App Service
        os.makedirs(output_dir, exist_ok=True)

        filename = f"{data.titulo}.{formato}"
        filepath = os.path.join(output_dir, filename)

        if formato == "txt":
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(f"Título: {data.titulo}\n")
                f.write(f"Autor: {data.autor}\n\n")
                f.write(data.contenido)
        else:
            doc = Document()
            doc.add_heading(data.titulo, 0)
            doc.add_paragraph(f"Autor: {data.autor}")
            doc.add_paragraph("")
            doc.add_paragraph(data.contenido)
            doc.save(filepath)

        return FileResponse(filepath, media_type="application/octet-stream", filename=filename)

    except Exception as e:
        logger.error(f"❌ Error al guardar artículo: {e}")
        raise HTTPException(status_code=500, detail=f"Error al guardar artículo: {str(e)}")
